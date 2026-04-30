from io import BytesIO

import fitz
from docx import Document


def extract_text_from_resume(file) -> str:
    """Extract text from PDF, DOCX, or TXT resume uploads."""
    name = getattr(file, "name", "").lower()
    content = file.getvalue() if hasattr(file, "getvalue") else file.read()

    if name.endswith(".pdf"):
        return extract_text_from_pdf_bytes(content)
    if name.endswith(".docx"):
        return extract_text_from_docx_bytes(content)
    if name.endswith(".txt"):
        return content.decode("utf-8", errors="ignore")

    raise ValueError("Unsupported resume format. Upload a PDF, DOCX, or TXT file.")


def extract_text_from_pdf(file) -> str:
    """Backward-compatible PDF extractor for existing callers."""
    content = file.getvalue() if hasattr(file, "getvalue") else file.read()
    return extract_text_from_pdf_bytes(content)


def extract_text_from_pdf_bytes(content: bytes) -> str:
    doc = fitz.open(stream=content, filetype="pdf")
    return "\n".join(page.get_text() for page in doc).strip()


def extract_text_from_docx_bytes(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    table_text = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_text.append(" | ".join(cells))

    return "\n".join(paragraphs + table_text).strip()
