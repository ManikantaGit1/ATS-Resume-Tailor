import io
import re
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def resume_markdown_to_docx(markdown_text: str, template_bytes: Optional[bytes] = None) -> bytes:
    """Convert generated resume Markdown into a downloadable DOCX."""
    document = _load_document(template_bytes)
    _clear_document_body(document)
    _configure_defaults(document)

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("# "):
            _add_name(document, line[2:].strip())
        elif line.startswith("## "):
            _add_section_heading(document, line[3:].strip())
        elif line.startswith("### "):
            _add_subheading(document, line[4:].strip())
        elif line.startswith(("- ", "* ")):
            _add_bullet(document, line[2:].strip())
        else:
            _add_paragraph(document, _strip_markdown(line))

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _load_document(template_bytes: Optional[bytes]) -> Document:
    if template_bytes:
        try:
            return Document(io.BytesIO(template_bytes))
        except Exception:
            pass
    return Document()


def _clear_document_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def _configure_defaults(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    for section in document.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)


def _add_name(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(_strip_markdown(text).upper())
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(20, 33, 61)


def _add_section_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(_strip_markdown(text).upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(20, 33, 61)
    _add_bottom_border(paragraph)


def _add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.22)
    paragraph.paragraph_format.first_line_indent = Inches(-0.14)
    paragraph.paragraph_format.space_after = Pt(2)
    bullet = paragraph.add_run("• ")
    bullet.font.size = Pt(10.5)
    bullet.font.color.rgb = RGBColor(20, 33, 61)

    _add_rich_text(paragraph, text, font_size=10.5)


def _add_subheading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(1)
    _add_rich_text(paragraph, f"**{text}**", font_size=10.5, color=RGBColor(20, 33, 61))


def _add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.05

    if _looks_like_contact_line(text):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rich_text = _role_line_markdown(text)
    _add_rich_text(paragraph, rich_text, font_size=10.5)


def _add_bottom_border(paragraph) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    paragraph_borders = paragraph_properties.find(qn("w:pBdr"))
    if paragraph_borders is None:
        paragraph_borders = OxmlElement("w:pBdr")
        paragraph_properties.append(paragraph_borders)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9DEE8")
    paragraph_borders.append(bottom)


def _strip_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    return text.strip()


def _add_rich_text(paragraph, text: str, font_size: float = 10.5, color=None) -> None:
    segments = re.split(r"(\*\*.*?\*\*)", text)
    for segment in segments:
        if not segment:
            continue
        is_bold = segment.startswith("**") and segment.endswith("**")
        clean = _strip_markdown(segment)
        if not clean:
            continue
        run = paragraph.add_run(clean)
        run.bold = is_bold
        run.font.size = Pt(font_size)
        if color is not None:
            run.font.color.rgb = color


def _role_line_markdown(text: str) -> str:
    clean = _strip_markdown(text)
    if "|" in clean:
        left, right = [part.strip() for part in clean.split("|", 1)]
        if left and right:
            return f"**{left}** | {right}"
    if re.search(r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\b.*\b(?:19|20)\d{2}\b", clean):
        return f"**{clean}**"
    return clean


def _looks_like_contact_line(text: str) -> bool:
    return bool(re.search(r"(@|linkedin|github|\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b)", text, re.I))
